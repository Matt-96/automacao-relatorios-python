from docx.shared import Inches

from docx import Document

from interface import *

from datetime import datetime
import os
#Essas funções manipulam arquivos .txt
def arquivo_existe(nome):
    try:
        a = open(nome, 'rt')
        a.close()
    except FileNotFoundError:
        return False
    else:
        return True

def criar_arquivo(nome):
        try:
            arq = open(nome, 'wt+')
            arq.close()
        except:
            print('Houve um erro na criação do arquivo')
        else:
            print(f'{nome} criado com sucesso')

def ler_arquivo(nome):
    try:
        arq = open(nome, 'rt')
    except:
        print('Erro ao ler o arquivo.')
    else:
        cabecalho('PESSOAS CADASTRADAS')
        for linha in arq:
            dados = linha.strip().split(';')
            if len(dados) > 1:
                print(f'{dados[0]:<30} {dados[1]:>3} anos')

def cadastro_pessoa(nome, pessoa='desconhecido', idade=0):
    try:
        #esse comando abre e fecha o arquivo
        with open(nome,'at') as a:
            #esse comando escreve os dados capturados pelo input do programa principal
            a.write(f'{pessoa};{idade}\n')
    except:
        print('Houve um ERRO ao escrever os dados!')
    else:
        print(f'Novo registro de {pessoa} adicionado.')

#Essas funções manipulam arquivos .docx

def gerar_docx(dados):
# Pega o nome do professor que está utilizando o programa
    #nome_professor = str(input('Informe o nome do professor: '))
    # Pergunta ao usuário se ele quer inserir a data manualmente ou que o sistema pegue a data atual
    escolha = validar('Deseja que o programa insira a data de hoje automaticamente? [S/N]')
    if escolha == 'S':
        # Pega a data atual
        data = datetime.today().strftime('%d-%m-%Y')
    else:
        data = str(input('Digite a data do relatório: ')).replace('/', '-')

# 1. Carrega o template (Se certifica que está na mesma pasta)
    try:
        doc = Document('RelatorioPCM.docx')
        tabela = doc.tables[0]


        # 2. Varre os parágrafos do topo do texto
        cabecalho_documento = tabela.cell(0,0)
        cabecalho_documento.text = f"RELATÓRIO SEMANAL DE ATIVIDADES      |    PROFESSOR(A): Matheus Felipe Dias Rocha   |    DATA: {data.replace('-', '/')}"

        # Loop que percorre e preenche as turmas
        # Range(1, 5) vai gerar os números de 1-4
        for i in range(2, 6):
            chave = f'T{i-1}'
            # Acessa linha i, coluna 1 e injeta o conteúdo formatado
            tabela.cell(i,1).text = formatar_conteudo(dados[chave])
        nome_arq = f'Relatorio de Violoncelo PCM{data}'


        #Insere a assinatura

        # --- ESTRATÉGIA DE SUBSTITUIÇÃO DIRETA ---
        assinatura_inserida = False

        for p in doc.paragraphs:
            if '[ASSINATURA_AQUI]' in p.text:
                # 1. Limpa o texto marcador
                p.text = p.text.replace('[ASSINATURA_AQUI]', '')

                # 2. Adiciona a imagem exatamente ali
                run = p.add_run()
                run.add_picture('assinatura.jpg', width=Inches(2))

                assinatura_inserida = True
                print("Assinatura inserida no marcador [ASSINATURA_AQUI].")
                break

        doc.save(f'Relatorio de Violoncelo PCM{data}.docx')
        salvar_no_drive(doc,f'Relatorio de Violoncelo PCM{data}.docx', data, r'G:\.shortcut-targets-by-id\1sVuR3CkYE0sF1k_7KeR0_tphFsJhXB4v\Crescendo com Música\Matheus')


    except Exception as erro:
        linha()
        print(f'ERRO! Não foi possível carregar o ficheiro. Erro: {erro} ')
        return
    else:
        linha()
        print(f'Ficheiro {nome_arq} criado com sucesso. ')


# 2. Mapeamento de Células (Linha, Coluna)
# Lógica de Formatação de Texto(vírgula, conectivo "e")

def formatar_conteudo(lista):
    if not lista:
        return 'Nenhum conteúdo registrado'
    if len(lista) == 1:
        return lista[0]
    return f'{','.join(lista[:-1])} e {lista[-1]}'


def salvar_no_drive(doc, nome_arquivo, data,diretorio_base):
    meses = {
        '01': 'Janeiro', '02': 'Fevereiro', '03': 'Março', '04': 'Abril',
        '05': 'Maio', '06': 'Junho', '07': 'Julho', '08': 'Agosto',
        '09': 'Setembro', '10': 'Outubro', '11': 'Novembro', '12': 'Dezembro'
    }

    try:
        # 1. Garante que a data seja tratada como texto e fatiada
        data_str = str(data)
        num_mes = data_str.split('-')[1]
        nome_mes = meses[num_mes]

        # 2. Caminho Base (Use o 'r' antes das aspas)
        caminho_base = r'G:\Meu Drive\RELATORIOS_TESTE_PCM'

        # 3. Criar a pasta do mês
        caminho_pasta_mes = os.path.join(diretorio_base, str(nome_mes))
        os.makedirs(caminho_pasta_mes, exist_ok=True)

        # 4. Salvar o arquivo
        caminho_final = os.path.join(caminho_pasta_mes, str(nome_arquivo))
        doc.save(caminho_final)

        print(f"Sucesso! Arquivo salvo em: {nome_mes}")

    except Exception as e:
        # Isso vai te mostrar exatamente onde está o erro se falhar de novo
        print(f"ERRO DETALHADO: {e}")
        import traceback
        traceback.print_exc()

def assinatura(doc, nome = 'Documento'):
    for p in doc.paragraphs:
        if '[ASSINATURA_AQUI]' in p.text:
            # 1. Limpa o texto marcador
            p.text = p.text.replace('[ASSINATURA_AQUI]', '')

            # 2. Adiciona a imagem exatamente ali
            run = p.add_run()
            run.add_picture('assinatura.jpg', width=Inches(2))

            assinatura_inserida = True
            print("Assinatura inserida no marcador [ASSINATURA_AQUI].")
            break
    doc.save(nome)


def gerar_relatorioJ(dados):
    # Pega o nome do professor que está utilizando o programa
    #nome_professor = str(input('Informe o nome do professor: '))
    # Pergunta ao usuário se ele quer inserir a data manualmente ou que o sistema pegue a data atual
    escolha = validar('Deseja que o programa insira a data de hoje automaticamente? [S/N]')
    if escolha == 'S':
        # Pega a data atual
        data = datetime.today().strftime('%d-%m-%Y')
    else:
        data = str(input('Digite a data do relatório: ')).replace('/', '-')

        # 1. Carrega o template (Se certifica que está na mesma pasta)
    try:
        doc = Document('Matheus - Juventude.docx')
        tabela = doc.tables[0]

        cabecalho_documento = tabela.cell(0, 0)
        cabecalho_documento.text = f"RELATÓRIO SEMANAL DE ATIVIDADES      |    PROFESSOR(A): Matheus Felipe Dias Rocha   |    DATA: {data.replace('-', '/')}"

        chave = 'T5'
        # Acessa linha i, coluna 1 e injeta o conteúdo formatado
        tabela.cell(2, 1).text = formatar_conteudo(dados[chave])
        nome_arq = f'Matheus - Juventude{data}'





    except Exception as erro:
        linha()
        print(f'ERRO! Não foi possível carregar o ficheiro. Erro: {erro} ')
        return
    else:
        linha()
        doc.save(f'Matheus - Juventude{data}.docx')
        assinatura(doc,f'Matheus - Juventude{data}.docx' )
        salvar_no_drive(doc,f'Matheus - Juventude{data}.docx',data,r'G:\.shortcut-targets-by-id\1YQ-CoCc-ciUu6ZB-64MmfDOOlACSV5pX\Juventude\Matheus')
        print(f'Ficheiro {nome_arq} criado com sucesso. ')